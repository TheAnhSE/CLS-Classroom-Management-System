"""
Updates Section 1.4.3 Non-UI Functions table in the CLS SRS .docx
Table index: 4  |  Header row: pink fill, bold
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCX = r'd:\Back up D and E\Work\Back up\NET\KI 8 _ Spring 2026\Block 5 _ SP2026\AI\Team Project\CLS-Classroom-Management-System\Document\02_Requirements\Report3_Software Requirement Specification.docx'

ROWS = [
    (1,  'Parent Notification Engine',
         'Attendance Notification Service',
         'Event-driven background service triggered immediately after attendance is recorded for any Session. '
         'Composes and dispatches an HTML email to the Parent via the Email Service API, including: session date/time, '
         'Teacher name, and Learner attendance status (Present / Absent / Late). '
         'Fully automated — zero manual action required from Academic Admin. (UC-19, BR-ATT-02)'),
    (2,  'Parent Notification Engine',
         'Schedule Change Notification Service',
         'Event-driven service triggered on any modification or cancellation of a Session '
         '(change to date, time, classroom, or Teacher). '
         'Broadcasts a notification email to all Parents of enrolled Learners, '
         'stating both original and updated schedule details. (UC-20, BR-NOT-01)'),
    (3,  'Parent Notification Engine',
         'Feedback Report Email Service',
         'Triggered automatically when a Teacher submits a Feedback entry for a Learner. '
         'Composes and dispatches a per-Learner email report to the associated Parent, '
         'containing academic rating (1–5), behavioral notes, and improvement recommendations. (UC-21)'),
    (4,  'Parent Notification Engine',
         'Package Expiry Cron Job',
         'Scheduled batch job running daily at 00:00 AM. '
         'Scans all active LearnerPackage records and identifies Learners whose remaining sessions '
         'reach the 2-week depletion threshold (BR-NOT-02). '
         'Dispatches expiry warning emails to associated Parents and creates in-app alerts for Academic Admin. '
         'Directly supports the 20% churn-reduction business goal. (UC-22, UC-23)'),
    (5,  'Session Scheduling',
         'Conflict Detection Algorithm',
         'Synchronous validation service invoked on every Session create or update (UC-13). '
         'Executes overlap query: WHERE (TimeStart <= RequestedEnd) AND (TimeEnd >= RequestedStart) '
         'for both RoomID and TeacherID. On conflict: blocks save and returns '
         'MSG-SCH-409 (Teacher double-booked) or MSG-SCH-410 (Room double-booked). '
         'Guarantees zero Scheduling Conflicts system-wide. (BR-SCH-01, BR-SCH-02)'),
    (6,  'Teacher Feedback',
         'SLA Compliance Tracker',
         'Background service triggered on each Feedback record insert. '
         'Calculates Delta = FeedbackSubmittedAt - SessionEndTime. '
         'Sets IsSLACompliant = True if Delta <= 12 hours; False otherwise. '
         'SLA data feeds into the Admin Operational Dashboard for breach monitoring and reporting. (UC-19, BR-FDB-02)'),
    (7,  'Teacher Feedback',
         'SLA Reminder Notification Service',
         'Scheduled monitoring service tracking pending Feedback entries for completed Sessions. '
         'Sends an in-app + email reminder to the Teacher at the 8-hour warning mark post-session, '
         'and a final reminder at the 12-hour SLA deadline. '
         'Late submissions are flagged automatically for Academic Admin oversight. (UC-18)'),
    (8,  'Learning Package Management',
         'Session Deduction Service',
         'Event-driven service triggered on each attendance record save (BR-ATT-01). '
         'Decrements the RemainingSession counter in the associated LearnerPackage by 1 '
         'for each Present or Late status. '
         'If RemainingSession reaches 0, logs an arrears warning: MSG-ATT-401. (UC-14)'),
    (9,  'System',
         'Activity Log Service',
         'Background service recording all significant user actions to the ActivityLog table: '
         'login/logout events, data creation/modification/deletion, report exports, and admin operations. '
         'Supports security monitoring, operational accountability, and full audit trail requirements.'),
    (10, 'System',
         'Notification Delivery Status Tracker',
         'Webhook receiver processing delivery status callbacks (delivered / bounced / failed) '
         'from the Email Service API (e.g., SendGrid, Amazon SES). '
         'Updates Notification log entries with delivery status and flags failed deliveries '
         'for Academic Admin review on the operational dashboard.'),
]

# ── XML helpers ────────────────────────────────────────────────────
def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def write_cell(cell, text, bold=False, size=10.5, align=None, color=None):
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if align:
        para.alignment = align

# ── Load & find table ──────────────────────────────────────────────
doc = Document(DOCX)
tbl = doc.tables[4]          # confirmed index from discovery run

# ── Rebuild table rows ─────────────────────────────────────────────
# 1) Clear existing rows except header (keep row 0)
while len(tbl.rows) > 1:
    tr = tbl.rows[-1]._tr
    tbl._tbl.remove(tr)

# 2) Update header row
hdr = tbl.rows[0]
headers = ['#', 'Feature', 'System Function', 'Description']
for i, cell in enumerate(hdr.cells):
    shade_cell(cell, 'F4CCCC')      # light pink header matching SRS template
    write_cell(cell, headers[i], bold=True, size=11,
               align=WD_ALIGN_PARAGRAPH.CENTER)

# 3) Add data rows
for num, feature, fn_name, desc in ROWS:
    row = tbl.add_row()
    cells = row.cells

    # Alternating fill: white / very light gray
    fill = 'FFFFFF' if num % 2 == 1 else 'F9F9F9'
    for c in cells:
        shade_cell(c, fill)

    write_cell(cells[0], str(num), align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5)
    write_cell(cells[1], feature, bold=False, size=10.5)
    write_cell(cells[2], fn_name, bold=True,  size=10.5)
    write_cell(cells[3], desc,    bold=False, size=10)

print(f'Table updated: {len(tbl.rows)} rows (1 header + {len(tbl.rows)-1} data rows)')

# ── Save ───────────────────────────────────────────────────────────
doc.save(DOCX)
print('Saved:', DOCX)
